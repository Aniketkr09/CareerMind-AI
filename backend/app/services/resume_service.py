import logging
import uuid
from pathlib import Path

from fastapi import HTTPException, UploadFile, status
from sqlalchemy.orm import Session

from app.models.resume import Resume


logger = logging.getLogger("CareerMindAI.ResumeService")


UPLOAD_DIR = Path("uploads/resumes")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)


ALLOWED_EXTENSIONS = {
    ".pdf",
    ".docx",
}


MAX_FILE_SIZE = 10 * 1024 * 1024


class ResumeService:

    @staticmethod
    async def upload_resume(
        db: Session,
        file: UploadFile,
        user_id,
    ) -> Resume:

        file_path: Path | None = None

        try:
            if file is None:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Resume file is required.",
                )

            if not file.filename:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Invalid resume filename.",
                )

            original_filename = Path(
                file.filename
            ).name

            extension = Path(
                original_filename
            ).suffix.lower()

            if extension not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Only PDF and DOCX files are allowed.",
                )

            content = await file.read()

            if not content:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail="Uploaded resume is empty.",
                )

            if len(content) > MAX_FILE_SIZE:
                raise HTTPException(
                    status_code=status.HTTP_413_REQUEST_ENTITY_TOO_LARGE,
                    detail="Resume file must be smaller than 10 MB.",
                )

            stored_filename = (
                f"{uuid.uuid4().hex}{extension}"
            )

            file_path = (
                UPLOAD_DIR / stored_filename
            )

            file_path.write_bytes(content)

            extracted_text = ResumeService.extract_text(
                str(file_path),
                extension,
            )

            if not extracted_text.strip():
                file_path.unlink(missing_ok=True)

                raise HTTPException(
                    status_code=status.HTTP_422_UNPROCESSABLE_ENTITY,
                    detail=(
                        "Unable to extract text from the resume. "
                        "Please upload a valid PDF or DOCX file."
                    ),
                )

            resume = Resume(
                user_id=user_id,
                original_filename=original_filename,
                stored_filename=stored_filename,
                file_path=str(file_path),
                file_type=extension,
                file_size=len(content),
                extracted_text=extracted_text,
                is_processed=False,
            )

            db.add(resume)
            db.commit()
            db.refresh(resume)

            logger.info(
                "Resume uploaded successfully | user=%s | resume=%s",
                user_id,
                resume.id,
            )

            return resume

        except HTTPException:
            db.rollback()

            if file_path and file_path.exists():
                file_path.unlink(missing_ok=True)

            raise

        except Exception as error:
            db.rollback()

            if file_path and file_path.exists():
                file_path.unlink(missing_ok=True)

            logger.exception(
                "Resume upload failed | user=%s | error=%s",
                user_id,
                error,
            )

            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to upload resume.",
            )

    @staticmethod
    def extract_text(
        file_path: str,
        extension: str,
    ) -> str:

        try:
            if extension == ".pdf":
                return ResumeService._extract_pdf_text(
                    file_path
                )

            if extension == ".docx":
                return ResumeService._extract_docx_text(
                    file_path
                )

            return ""

        except Exception as error:
            logger.exception(
                "Resume text extraction failed | file=%s | error=%s",
                file_path,
                error,
            )

            return ""

    @staticmethod
    def _extract_pdf_text(
        file_path: str,
    ) -> str:

        import PyPDF2

        text_parts: list[str] = []

        with open(
            file_path,
            "rb",
        ) as file:

            reader = PyPDF2.PdfReader(file)

            for page in reader.pages:
                page_text = (
                    page.extract_text()
                    or ""
                )

                if page_text.strip():
                    text_parts.append(
                        page_text.strip()
                    )

        return "\n".join(
            text_parts
        ).strip()

    @staticmethod
    def _extract_docx_text(
        file_path: str,
    ) -> str:

        from docx import Document

        document = Document(file_path)

        text_parts: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()

            if text:
                text_parts.append(text)

        for table in document.tables:
            for row in table.rows:

                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]

                if cells:
                    text_parts.append(
                        " | ".join(cells)
                    )

        return "\n".join(
            text_parts
        ).strip()

    @staticmethod
    def get_user_resumes(
        db: Session,
        user_id,
    ) -> list[Resume]:

        return (
            db.query(Resume)
            .filter(
                Resume.user_id == user_id
            )
            .order_by(
                Resume.created_at.desc()
            )
            .all()
        )

    @staticmethod
    def get_resume(
        db: Session,
        resume_id,
        user_id,
    ) -> Resume:

        resume = (
            db.query(Resume)
            .filter(
                Resume.id == resume_id,
                Resume.user_id == user_id,
            )
            .first()
        )

        if not resume:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="Resume not found.",
            )

        return resume

    @staticmethod
    def get_latest_resume(
        db: Session,
        user_id,
    ) -> Resume:

        resume = (
            db.query(Resume)
            .filter(
                Resume.user_id == user_id
            )
            .order_by(
                Resume.created_at.desc()
            )
            .first()
        )

        if not resume:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="No resume found.",
            )

        return resume

    @staticmethod
    def delete_resume(
        db: Session,
        resume: Resume,
    ) -> dict:

        try:
            file_path = Path(
                resume.file_path
            )

            if file_path.exists():
                file_path.unlink()

            db.delete(resume)
            db.commit()

            logger.info(
                "Resume deleted successfully | resume=%s",
                resume.id,
            )

            return {
                "message": "Resume deleted successfully."
            }

        except Exception as error:
            db.rollback()

            logger.exception(
                "Resume deletion failed | resume=%s | error=%s",
                resume.id,
                error,
            )

            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail="Failed to delete resume.",
            )